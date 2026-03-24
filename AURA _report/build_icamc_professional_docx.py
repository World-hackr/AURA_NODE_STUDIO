from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "ICAMC2026_AURA_Conference_Professional_Rewrite.docx"
OUTPUT_MD = ROOT / "ICAMC2026_AURA_Conference_Professional_Rewrite.md"

FIG_PROBLEM = ROOT / "generated_aura_problem_goal.png"
FIG_OVERVIEW = ROOT / "generated_aura_system_overview.png"
FIG_WORKFLOW = ROOT / "generated_aura_workflow.png"
FIG_SURFACES = ROOT / "generated_aura_surfaces.png"

TITLE = (
    "AURA: A Complete Cyber-Physical Prototype for Inventory-Aware "
    "Circuit Creation and Physical Component Retrieval"
)

AUTHOR_LINES = [
    "Santosh Kumar",
    "Department of Computer Science and Engineering",
    "HMR Institute of Technology and Management, New Delhi, India",
    "Academic guidance: Dr. Naveen Sharma, Professor",
    "Supervisor: Dr. Naveen Sharma, Professor",
]

ABSTRACT = (
    "Electronics workbench software and physical storage systems are usually disconnected. "
    "When a user wants to build a circuit, two bottlenecks appear immediately: determining "
    "what can be created from the exact parts already owned, and physically retrieving "
    "those parts from drawers, bins, or racks once the design has been chosen. This paper "
    "presents AURA, a complete cyber-physical prototype that addresses both bottlenecks "
    "within one inventory-aware system. AURA combines structured AI-assisted circuit "
    "creation, exact-part inventory records, an ESP32-based host with a compact screen and "
    "joystick interface, and distributed locator nodes built around the ATtiny404, "
    "nRF24L01, and WS2812-based indication. On the design side, AURA grounds candidate "
    "circuits in owned-stock awareness and structured review rather than unconstrained chat "
    "output. On the physical side, the host resolves exact storage locations and activates "
    "addressed nodes to highlight stored components for retrieval. The current prototype "
    "direction supports inventory-constrained circuit assistance, single-part lookup, "
    "host-side stock review, node-addressed localization, and compact local control "
    "without requiring a phone during retrieval. The goal of AURA is to reduce both "
    "design-side uncertainty and search-side friction, so that a user can move more "
    "directly from build intent to physically assembled hardware. The paper frames AURA as "
    "a complete prototype and argues that its core practical contribution lies in unifying "
    "circuit creation, inventory fit, and real-world component retrieval within one "
    "inspectable workflow."
)

INDEX_TERMS = (
    "cyber-physical systems, electronics inventory, AI-assisted circuit creation, "
    "component retrieval, distributed locator nodes, inventory-aware workflow"
)

SECTIONS = []

REFERENCES = [
    '[1] Espressif Systems, "ESP32 Technical Reference Manual," 2023.',
    '[2] Microchip Technology Inc., "tinyAVR 0-series Family Data Sheet," 2024.',
    '[3] Nordic Semiconductor, "nRF24L01+ Single Chip 2.4 GHz Transceiver Product Specification," 2019.',
    '[4] Worldsemi, "WS2812B Intelligent Control LED Integrated Light Source Data Sheet," 2020.',
    '[5] Bluetooth SIG, "Bluetooth Core Specification," Version 5.x.',
    '[6] Arduino, "SPI Communication," Arduino Documentation.',
    '[7] SQLite, "SQLite Documentation," Documentation.',
    '[8] P. Horowitz and W. Hill, The Art of Electronics, 3rd ed. Cambridge, U.K.: Cambridge Univ. Press, 2015.',
    '[9] IETF, "JavaScript Object Notation (JSON) Patch," RFC 6902, 2013.',
]


def font(size, bold=False):
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path("C:/Windows/Fonts") / font_name
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_12 = font(12)
FONT_14B = font(14, True)
FONT_16B = font(16, True)
FONT_18B = font(18, True)
FONT_22B = font(22, True)


SECTIONS = [
    {
        "title": "I. INTRODUCTION",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "Electronics assistance usually breaks down exactly where real work "
                    "begins. A user may know the function they want, such as a timer, a "
                    "flasher, a motor driver, or a simple control circuit, but still face "
                    "two immediate barriers. First, it is unclear what can be built from "
                    "the exact parts already owned. Second, even after the required parts "
                    "are identified, it is still unclear where those parts are physically "
                    "stored and how quickly they can be gathered for assembly."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "This double bottleneck is more serious than it first appears. Circuit "
                    "tools often assume ideal availability. Inventory tools often stop at "
                    "record keeping. General-purpose AI may produce plausible circuit ideas, "
                    "but usually does not remain grounded in exact stock, exact quantities, "
                    "and physical storage reality. As a result, users are forced to jump "
                    "between design thinking, stock checking, and manual search inside the "
                    "real storage environment."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "AURA is proposed as a direct answer to this fragmentation. Its goal is "
                    "not only to help the user think about a circuit, and not only to help "
                    "the user find a part. Its goal is to connect both tasks through one "
                    "inventory-aware cyber-physical workflow so that build intent can become "
                    "a reviewable circuit direction, an accepted part set, and finally a "
                    "successful real-world retrieval path."
                ),
            },
            {
                "type": "figure",
                "path": FIG_PROBLEM,
                "caption": (
                    "Fig. 1. AURA targets two linked bottlenecks at once: inventory-aware "
                    "circuit creation and physical component retrieval."
                ),
                "width_in": 3.2,
            },
            {
                "type": "subheading",
                "text": "A. Paper Goal",
            },
            {
                "type": "paragraph",
                "text": (
                    "The central goal of the paper is to make the reader understand AURA's "
                    "purpose without ambiguity. AURA is a hardware-first electronics "
                    "assistance system that helps users 1) decide what can be built from "
                    "owned stock, 2) verify whether enough quantity exists, 3) map accepted "
                    "needs to exact storage locations, and 4) retrieve those parts through "
                    "a host-and-node localization system."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Main Contributions",
            },
            {
                "type": "paragraph",
                "text": (
                    "The main contribution of AURA is the integration itself. The system "
                    "brings together an inventory-aware circuit-creation path, a compact "
                    "embedded host, distributed locator nodes, and exact-part storage "
                    "mapping within one complete prototype direction. More specifically, the "
                    "paper contributes 1) a structured circuit-assistance layer grounded in "
                    "owned inventory, 2) a physical retrieval layer that highlights exact "
                    "storage positions through addressed nodes, 3) a shared inventory model "
                    "for exact values, quantities, and locations, and 4) an end-to-end "
                    "workflow from build intent to physically retrieved parts."
                ),
            },
        ],
    },
    {
        "title": "II. SYSTEM GOALS AND OVERVIEW",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "AURA should be understood as a system with clearly bounded goals. It "
                    "must reduce design-side uncertainty, reduce search-side friction, keep "
                    "AI assistance grounded in structured inventory knowledge, and preserve "
                    "a practical host-first retrieval experience. These goals define the "
                    "product more accurately than labels such as editor, simulator, or "
                    "inventory app."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "The system is built around one shared inventory substrate. That "
                    "inventory layer holds exact component identity, value or variant, "
                    "quantity, storage address, and node-output mapping. From that common "
                    "substrate, one branch supports circuit-oriented assistance and another "
                    "branch supports host-guided localization. This shared base is what "
                    "keeps AURA from becoming two disconnected tools."
                ),
            },
            {
                "type": "figure",
                "path": FIG_OVERVIEW,
                "caption": (
                    "Fig. 2. Professional overview of the current AURA system direction, "
                    "showing how shared inventory supports both circuit creation and "
                    "physical retrieval."
                ),
                "width_in": 3.2,
            },
            {
                "type": "subheading",
                "text": "A. Circuit-Creation Goal",
            },
            {
                "type": "paragraph",
                "text": (
                    "On the design side, AURA should help the user move from intent toward "
                    "a candidate circuit that is not only plausible, but grounded in the "
                    "parts actually available. This means the system must reason over exact "
                    "values and quantities rather than abstract families alone. AURA is not "
                    "intended to replace engineering judgment or claim universal electrical "
                    "correctness. Instead, it should provide structured assistance that "
                    "keeps the design state inspectable and constrained."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Retrieval Goal",
            },
            {
                "type": "paragraph",
                "text": (
                    "On the physical side, AURA should do more than report that a part is "
                    "owned somewhere. It should resolve the exact storage location and "
                    "activate an addressed output that guides the user to the correct drawer, "
                    "strip, or storage zone. This makes retrieval a first-class system "
                    "function instead of an external manual step."
                ),
            },
            {
                "type": "subheading",
                "text": "C. Host-First Goal with Software Support",
            },
            {
                "type": "paragraph",
                "text": (
                    "AURA does include phone and software support, but these remain support "
                    "surfaces rather than the primary product identity. The host device must "
                    "remain sufficient for local lookup, stock review, and locate actions. "
                    "Richer software may assist with bulk input, project organization, and "
                    "broader search, but the system should still function as a real hardware "
                    "product even when the user is standing near storage with only the host "
                    "in hand."
                ),
            },
        ],
    },
    {
        "title": "III. COMPLETE PROTOTYPE ARCHITECTURE",
        "blocks": [
            {
                "type": "subheading",
                "text": "A. Inventory-Aware Circuit Assistance Layer",
            },
            {
                "type": "paragraph",
                "text": (
                    "The circuit-assistance layer accepts build intent and inventory context "
                    "and returns structured candidate results rather than opaque free-form "
                    "responses. Its purpose is not to present AI as unquestionable authority, "
                    "but to keep proposals reviewable, comparable, and tied to owned stock. "
                    "This is the design-side half of AURA's contribution."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Host Device Layer",
            },
            {
                "type": "paragraph",
                "text": (
                    "The host prototype is based on an ESP32-class controller with a 1.8 inch "
                    "local display and joystick-driven input. It acts as the embedded control "
                    "surface for part lookup, stock review, node test, setup actions, and "
                    "live locate sessions. The host should remain understandable and useful "
                    "without long text entry or complex on-device editing."
                ),
            },
            {
                "type": "figure",
                "path": ROOT / "ppt_assets" / "conference_media" / "image13.png",
                "caption": (
                    "Fig. 3. Current AURA host prototype with compact display and joystick-"
                    "driven local interaction."
                ),
                "width_in": 3.0,
            },
            {
                "type": "subheading",
                "text": "C. Distributed Locator Node Layer",
            },
            {
                "type": "paragraph",
                "text": (
                    "The node layer uses compact low-power units built around the ATtiny404 "
                    "and nRF24L01, combined with WS2812-based visual indication. A node can "
                    "cover one dense row, one rack segment, or another mapped physical zone. "
                    "This layer provides the actuation path that turns storage metadata into "
                    "visible retrieval guidance."
                ),
            },
            {
                "type": "figure",
                "path": ROOT / "ppt_assets" / "conference_media" / "image14.png",
                "caption": (
                    "Fig. 4. Current locator node prototype for addressed visual highlighting "
                    "of stored components."
                ),
                "width_in": 2.6,
            },
            {
                "type": "subheading",
                "text": "D. Current Interaction Surfaces",
            },
            {
                "type": "paragraph",
                "text": (
                    "The current AURA direction intentionally separates rich support "
                    "surfaces from compact retrieval surfaces. The phone or support app "
                    "handles deeper search, inventory intake, project grouping, and richer "
                    "review tasks. The host focuses on fast local control. This split is "
                    "important because it lets the paper present a realistic cyber-physical "
                    "system instead of pretending that all tasks belong equally on the same "
                    "tiny display."
                ),
            },
            {
                "type": "figure",
                "path": FIG_SURFACES,
                "caption": (
                    "Fig. 5. Conceptual support surfaces aligned to the current AURA restart: "
                    "phone-side deep workflows and host-side retrieval control."
                ),
                "width_in": 3.2,
            },
        ],
    },
    {
        "title": "IV. END-TO-END WORKFLOW AND EFFECT",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "AURA is most meaningful when read as one continuous operational flow. "
                    "The user begins with build intent or with a known part need. The system "
                    "then moves through inventory-aware circuit assistance, exact-part "
                    "resolution, stock fit, accepted part selection, and finally physical "
                    "retrieval through addressed highlighting."
                ),
            },
            {
                "type": "figure",
                "path": FIG_WORKFLOW,
                "caption": (
                    "Fig. 6. End-to-end AURA workflow from build intent to inventory-aware "
                    "decision and finally to physical retrieval."
                ),
                "width_in": 3.2,
            },
            {
                "type": "subheading",
                "text": "A. Build Intent to Candidate Circuit",
            },
            {
                "type": "paragraph",
                "text": (
                    "The first effect of AURA is that it shortens the path from vague intent "
                    "to a concrete candidate circuit. Instead of forcing the user to start "
                    "with a theoretically ideal design and then discover missing parts later, "
                    "AURA keeps the early design process linked to what is actually available."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Candidate Circuit to Exact Part Set",
            },
            {
                "type": "paragraph",
                "text": (
                    "Once a candidate circuit is accepted, AURA turns that structured result "
                    "into a required part set grounded in exact values and quantities. This "
                    "stage is critical because it bridges the software-side decision with the "
                    "physical-side action. Without this step, design assistance and retrieval "
                    "would remain disconnected."
                ),
            },
            {
                "type": "subheading",
                "text": "C. Exact Part Set to Physical Retrieval",
            },
            {
                "type": "paragraph",
                "text": (
                    "After the required part set is known, the host resolves storage "
                    "addresses and activates the mapped locator outputs. The user is then "
                    "guided toward the correct storage position rather than performing a "
                    "manual room-scale search. This is where AURA's cyber-physical nature "
                    "becomes most visible."
                ),
            },
            {
                "type": "subheading",
                "text": "D. Practical User Effect",
            },
            {
                "type": "paragraph",
                "text": (
                    "The practical effect is that AURA reduces two forms of friction at once. "
                    "It reduces design-side uncertainty by grounding circuit assistance in "
                    "owned stock, and it reduces search-side friction by making physical "
                    "retrieval explicit and guided. This dual reduction is the strongest "
                    "reader-facing argument for why the system matters."
                ),
            },
        ],
    },
    {
        "title": "V. CURRENT PROTOTYPE STATUS AND VALIDATION DIRECTION",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "The paper frames AURA as a complete prototype direction because the "
                    "system is defined as one integrated product rather than as disconnected "
                    "ideas. The host, locator node concept, inventory model, workflow, and "
                    "interaction surfaces already form one coherent operational target. In "
                    "that sense, the work is not positioned as an isolated algorithm or a "
                    "detached UI concept, but as a prototype system whose remaining effort "
                    "is primarily refinement, implementation completion, and validation."
                ),
            },
            {
                "type": "subheading",
                "text": "A. Current Prototype Scope",
            },
            {
                "type": "paragraph",
                "text": (
                    "The current scope includes host-guided part lookup, stock-aware "
                    "reasoning, mapped node-based localization, compact local UI design, and "
                    "support-surface planning for richer workflows. It also includes a "
                    "clear contract that circuit assistance should remain inventory-aware and "
                    "physically actionable rather than floating as a separate virtual tool."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Validation Direction",
            },
            {
                "type": "paragraph",
                "text": (
                    "Prototype validation should focus on realistic operational scenarios: "
                    "single-part retrieval, inventory-constrained candidate-circuit review, "
                    "accepted-part-list to locate transition, host-only operation near "
                    "storage, and node mapping or remapping tasks. These validation paths "
                    "matter more to AURA than traditional software-only benchmark claims "
                    "because the system's value is inherently cyber-physical."
                ),
            },
            {
                "type": "subheading",
                "text": "C. Claim Discipline",
            },
            {
                "type": "paragraph",
                "text": (
                    "The paper should speak confidently about the integration and the product "
                    "goal, but it should avoid casual global claims such as 'first in the "
                    "world' unless they are explicitly defended through prior-art analysis. "
                    "The safer and stronger position is that AURA offers a rare and highly "
                    "practical unification of inventory-aware circuit creation and physical "
                    "component retrieval."
                ),
            },
        ],
    },
    {
        "title": "VI. CONCLUSION",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "AURA is best understood as a complete cyber-physical prototype whose "
                    "problem statement is simple and practical: help the user decide what "
                    "can be built from owned parts and then help the user retrieve those "
                    "parts in physical space. By joining structured circuit assistance, "
                    "inventory awareness, host-guided lookup, and node-based highlighting, "
                    "the system turns what are normally separate activities into one "
                    "inspectable workflow."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "This is the paper's central message and should remain visible from the "
                    "title through the conclusion. AURA is not merely a circuit editor, not "
                    "merely an inventory tool, and not merely a locator network. It is a "
                    "unified electronics assistance system that connects design intent, "
                    "inventory reality, and physical retrieval."
                ),
            },
        ],
    },
]


def wrap_text(draw, text, use_font, max_width):
    words = text.split()
    lines = []
    current = []
    for word in words:
        trial = " ".join(current + [word])
        if current and draw.textlength(trial, font=use_font) > max_width:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


def draw_text_block(draw, xy, text, use_font, fill, max_width, line_height):
    x, y = xy
    for line in wrap_text(draw, text, use_font, max_width):
        draw.text((x, y), line, font=use_font, fill=fill)
        y += line_height
    return y


def draw_card(draw, rect, title, body, accent):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=24, fill="white", outline=(40, 40, 40), width=3)
    draw.rounded_rectangle((x1, y1, x2, y1 + 34), radius=24, fill=accent, outline=accent)
    draw.rectangle((x1, y1 + 18, x2, y1 + 34), fill=accent, outline=accent)
    draw.text((x1 + 16, y1 + 7), title, font=FONT_16B, fill="white")
    draw_text_block(draw, (x1 + 16, y1 + 52), body, FONT_12, (20, 20, 20), x2 - x1 - 32, 18)


def draw_arrow(draw, start, end, fill=(50, 50, 50), width=4):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    if abs(x2 - x1) > abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        draw.polygon(
            [(x2, y2), (x2 - 14 * direction, y2 - 7), (x2 - 14 * direction, y2 + 7)],
            fill=fill,
        )
    else:
        direction = 1 if y2 > y1 else -1
        draw.polygon(
            [(x2, y2), (x2 - 7, y2 - 14 * direction), (x2 + 7, y2 - 14 * direction)],
            fill=fill,
        )


def create_problem_goal_figure():
    img = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "AURA Problem and Goal", font=FONT_22B, fill=(20, 20, 20))
    draw_text_block(
        draw,
        (60, 90),
        "AURA is designed around a clear reader-facing problem: users need help deciding what "
        "can be built from owned parts and then physically retrieving those parts.",
        FONT_12,
        (60, 60, 60),
        1450,
        18,
    )
    draw_card(
        draw,
        (70, 180, 760, 520),
        "Design-Side Bottleneck",
        "Users may know the function they want, but do not know what circuit can be created "
        "from the exact parts, values, and quantities they already own.",
        (34, 64, 94),
    )
    draw_card(
        draw,
        (840, 180, 1530, 520),
        "Retrieval-Side Bottleneck",
        "Even after the needed parts are identified, users still lose time searching drawers, "
        "bins, and racks because physical retrieval is usually not part of the tool itself.",
        (70, 102, 62),
    )
    draw_arrow(draw, (760, 350), (840, 350))
    draw.rounded_rectangle((220, 610, 1380, 900), radius=30, fill=(245, 247, 250), outline=(30, 30, 30), width=3)
    draw.text((260, 650), "AURA Goal", font=FONT_18B, fill=(20, 20, 20))
    draw_text_block(
        draw,
        (260, 700),
        "Turn build intent into a reviewable, inventory-aware circuit direction and then "
        "turn accepted parts into physically retrievable storage targets through a host-and-node system.",
        FONT_14B,
        (20, 20, 20),
        1080,
        22,
    )
    img.save(FIG_PROBLEM)


def create_system_overview_figure():
    img = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "AURA System Overview", font=FONT_22B, fill=(20, 20, 20))
    draw_card(
        draw,
        (610, 120, 1090, 270),
        "Shared Inventory Substrate",
        "Exact part identity, value or variant, quantity, storage address, node mapping, and project linkage.",
        (55, 88, 124),
    )
    draw_card(
        draw,
        (80, 380, 590, 640),
        "Inventory-Aware Circuit Creation",
        "Build intent, structured assistance, candidate circuits, review, accepted part list.",
        (79, 107, 56),
    )
    draw_card(
        draw,
        (620, 380, 1080, 640),
        "AURA Host",
        "Local lookup, stock review, compact UI, node test, locate activation, host-side control.",
        (38, 70, 93),
    )
    draw_card(
        draw,
        (1110, 380, 1620, 640),
        "Distributed Locator Nodes",
        "Addressed indication, mapped outputs, physical highlighting, storage guidance.",
        (112, 79, 36),
    )
    draw_card(
        draw,
        (200, 760, 610, 940),
        "Phone / Support App",
        "Deep search, bulk intake, project grouping, notes, richer setup support.",
        (95, 95, 95),
    )
    draw_card(
        draw,
        (760, 760, 1460, 940),
        "Real Storage Environment",
        "Drawers, bins, racks, shelves, project boxes, and mapped physical retrieval zones.",
        (95, 95, 95),
    )
    draw_arrow(draw, (850, 270), (335, 380))
    draw_arrow(draw, (850, 270), (850, 380))
    draw_arrow(draw, (850, 270), (1365, 380))
    draw_arrow(draw, (590, 510), (620, 510))
    draw_arrow(draw, (1080, 510), (1110, 510))
    draw_arrow(draw, (850, 640), (850, 760))
    draw_arrow(draw, (405, 760), (405, 640))
    draw.text((690, 700), "accepted part set becomes physical locate target", font=FONT_12, fill=(60, 60, 60))
    img.save(FIG_OVERVIEW)


def create_workflow_figure():
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "AURA End-to-End Workflow", font=FONT_22B, fill=(20, 20, 20))
    steps = [
        ("Build Intent", "User starts from function, need, or target behavior."),
        ("Inventory-Aware Assistance", "System proposes or refines circuit directions grounded in owned stock."),
        ("Accepted Part Set", "Exact values, quantities, and storage mappings are resolved."),
        ("Host Locate Action", "Host shows stock and activates mapped node targets."),
        ("Physical Retrieval", "User retrieves exact parts from highlighted storage positions."),
    ]
    x = 70
    y1 = 180
    w = 310
    h = 320
    accents = [(43, 70, 95), (66, 110, 78), (95, 88, 50), (84, 70, 116), (100, 65, 65)]
    centers = []
    for idx, (title, body) in enumerate(steps):
        rect = (x, y1, x + w, y1 + h)
        draw_card(draw, rect, title, body, accents[idx])
        centers.append((x + w, y1 + h / 2))
        x += 345
    for i in range(len(centers) - 1):
        draw_arrow(draw, (centers[i][0], centers[i][1]), (centers[i + 1][0] - 35, centers[i + 1][1]))
    draw_text_block(
        draw,
        (120, 560),
        "Effect: design-side uncertainty is reduced first, then search-side friction is reduced in physical space.",
        FONT_14B,
        (30, 30, 30),
        1450,
        22,
    )
    img.save(FIG_WORKFLOW)


def create_surfaces_figure():
    img = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "Current AURA Interaction Surfaces", font=FONT_22B, fill=(20, 20, 20))
    draw_text_block(
        draw,
        (60, 90),
        "These are conceptual surfaces aligned to the current AURA restart, not reused artwork from older software-first versions.",
        FONT_12,
        (60, 60, 60),
        1500,
        18,
    )
    draw.rounded_rectangle((80, 180, 720, 940), radius=38, fill=(248, 248, 248), outline=(30, 30, 30), width=4)
    draw.rounded_rectangle((115, 230, 685, 900), radius=24, fill=(255, 255, 255), outline=(50, 50, 50), width=2)
    draw.text((270, 190), "Phone / Support App", font=FONT_18B, fill=(20, 20, 20))
    phone_cards = [
        ("Search", "Deep part search and exact-value browsing."),
        ("Inventory", "Bulk intake, quantities, notes, images."),
        ("Projects", "Accepted parts grouped by build or task."),
        ("Nodes", "Setup assistance and mapping support."),
    ]
    px, py = 145, 265
    for title, body in phone_cards:
        draw_card(draw, (px, py, 655, py + 120), title, body, (52, 79, 104))
        py += 145

    draw.rounded_rectangle((930, 180, 1540, 940), radius=40, fill=(250, 250, 250), outline=(30, 30, 30), width=4)
    draw.rounded_rectangle((1040, 240, 1430, 520), radius=18, fill=(245, 245, 245), outline=(40, 40, 40), width=3)
    draw.text((1100, 190), "AURA Host Remote", font=FONT_18B, fill=(20, 20, 20))
    screens = [
        ("HOME", ["> Locate Parts", "  Inventory", "  Nodes", "  Setup"]),
        ("SEARCH", ["Query: 220R", "> 220R resistor", "  220R pack", "  220R 1/8W"]),
        ("SUMMARY", ["Part: 220R", "Stock: 43", "Place: Drawer A3", "> Start Locate"]),
        ("LOCATE", ["Node: N02", "Output: 14", "Place: Drawer A3", "> Stop Locate"]),
    ]
    sx, sy = 970, 560
    for idx, (title, lines) in enumerate(screens):
        x1 = sx + (idx % 2) * 230
        y1 = sy + (idx // 2) * 150
        draw.rounded_rectangle((x1, y1, x1 + 200, y1 + 120), radius=16, fill="white", outline=(40, 40, 40), width=2)
        draw.rectangle((x1 + 10, y1 + 10, x1 + 190, y1 + 28), fill=(40, 40, 40))
        draw.text((x1 + 18, y1 + 12), title, font=FONT_12, fill="white")
        ly = y1 + 40
        for line in lines:
            draw.text((x1 + 16, ly), line, font=FONT_12, fill=(20, 20, 20))
            ly += 18

    draw.ellipse((1160, 350, 1310, 500), outline=(50, 50, 50), width=3)
    draw.ellipse((1215, 405, 1255, 445), fill=(220, 220, 220), outline=(50, 50, 50))
    draw.rectangle((1115, 600, 1135, 630), fill=(90, 90, 90))
    draw.rectangle((1138, 600, 1158, 630), fill=(90, 90, 90))
    draw.text((1100, 530), "1.8 inch LCD", font=FONT_12, fill=(30, 30, 30))
    draw.text((1160, 515), "Joystick", font=FONT_12, fill=(30, 30, 30))
    draw.text((1065, 640), "Home / Back", font=FONT_12, fill=(30, 30, 30))
    img.save(FIG_SURFACES)


def generate_figures():
    create_problem_goal_figure()
    create_system_overview_figure()
    create_workflow_figure()
    create_surfaces_figure()


def set_page_layout(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.63)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def set_two_columns(section, num=2, space_twips=360):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    if cols.getparent() is None:
        sect_pr.append(cols)


def set_run_font(run, size_pt, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph_text(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.16):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Inches(first_indent)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size)
    return p


def add_center_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 10, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 10, bold=True)
    return p


def add_abstract_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(f"{label}\u2014")
    set_run_font(label_run, 9, bold=True, italic=True)
    body_run = p.add_run(text)
    set_run_font(body_run, 9)
    return p


def add_figure(doc, path, caption, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    set_run_font(run, 8, italic=True)


def add_references(doc):
    add_center_heading(doc, "REFERENCES")
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(ref)
        set_run_font(run, 8)


def build_markdown():
    lines = [
        f"# {TITLE}",
        "",
        *AUTHOR_LINES,
        "",
        "## Abstract",
        "",
        ABSTRACT,
        "",
        "## Index Terms",
        "",
        INDEX_TERMS,
        "",
    ]
    for section in SECTIONS:
        lines.append(f"## {section['title']}")
        lines.append("")
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                lines.append(block["text"])
                lines.append("")
            elif block["type"] == "subheading":
                lines.append(f"### {block['text']}")
                lines.append("")
            elif block["type"] == "figure":
                lines.append(f"![{block['caption']}]({block['path'].as_posix()})")
                lines.append("")
                lines.append(f"*{block['caption']}*")
                lines.append("")
    lines.append("## REFERENCES")
    lines.append("")
    lines.extend(REFERENCES)
    lines.append("")
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_docx():
    doc = Document()
    set_page_layout(doc.sections[0])
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(TITLE)
    set_run_font(run, 18, bold=True)

    for line in AUTHOR_LINES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, 10)

    doc.add_paragraph()
    add_abstract_paragraph(doc, "Abstract", ABSTRACT)
    add_abstract_paragraph(doc, "Index Terms", INDEX_TERMS)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page_layout(body_section)
    set_two_columns(body_section)

    for section in SECTIONS:
        add_center_heading(doc, section["title"])
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                add_paragraph_text(doc, block["text"])
            elif block["type"] == "subheading":
                add_subheading(doc, block["text"])
            elif block["type"] == "figure":
                add_figure(doc, block["path"], block["caption"], block["width_in"])

    add_references(doc)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    generate_figures()
    build_markdown()
    build_docx()
    print(f"Wrote {FIG_PROBLEM.name}")
    print(f"Wrote {FIG_OVERVIEW.name}")
    print(f"Wrote {FIG_WORKFLOW.name}")
    print(f"Wrote {FIG_SURFACES.name}")
    print(f"Wrote {OUTPUT_MD.name}")
    print(f"Wrote {OUTPUT_DOCX.name}")
