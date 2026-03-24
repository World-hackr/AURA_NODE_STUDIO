from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "ICAMC2026_AURA_Conference_Balanced_Rewrite.docx"
OUTPUT_MD = ROOT / "ICAMC2026_AURA_Conference_Balanced_Rewrite.md"
ARCHITECTURE_FIG = ROOT / "generated_aura_dual_architecture.png"


TITLE = (
    "AURA: A Complete Cyber-Physical Prototype for Inventory-Aware "
    "AI-Assisted Circuit Creation and Distributed Component Localization"
)

AUTHOR_LINES = [
    "Santosh Kumar",
    "Department of Computer Science and Engineering",
    "HMR Institute of Technology and Management, New Delhi, India",
    "Academic guidance: Dr. Naveen Sharma, Professor",
    "Supervisor: Dr. Naveen Sharma, Professor",
]

ABSTRACT = (
    "Electronics assistance usually splits into two separate problems: deciding what "
    "circuit can be built from available parts and physically finding those parts in "
    "storage. This paper presents AURA, a complete cyber-physical prototype that "
    "addresses both problems within one integrated system. AURA combines value-aware "
    "inventory records, a deterministic structured AI-assisted circuit creation layer, "
    "an ESP32-based host device with local interface, and distributed ATtiny404 plus "
    "nRF24L01 locator nodes with WS2812-based indication. On the design side, the "
    "system constrains circuit creation through owned-stock awareness and structured, "
    "reviewable outputs instead of unconstrained chat responses. On the physical side, "
    "the host resolves exact part locations and activates addressed nodes to highlight "
    "stored components for retrieval. The current prototype supports single-part lookup, "
    "inventory-constrained circuit proposals, structured review of candidate builds, "
    "and node-based localization of accepted parts. By coupling circuit creation and "
    "physical retrieval through the same inventory substrate, AURA reduces both design-"
    "side and retrieval-side friction while preserving a clear boundary between "
    "deterministic logic, assistive AI, and real-world hardware actions."
)

INDEX_TERMS = (
    "cyber-physical systems, AI-assisted circuit creation, electronics inventory, "
    "component localization, deterministic circuit assistance, inventory-aware retrieval"
)

SECTIONS = [
    {
        "title": "I. INTRODUCTION",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "Electronics prototyping often fails at two connected stages. The first "
                    "is design feasibility: users may know the outcome they want, but do "
                    "not know what circuit can be realistically created from the exact parts "
                    "they already own. The second is physical retrieval: even after the "
                    "needed components are identified, users still lose time checking "
                    "whether those exact values exist in sufficient quantity and where they "
                    "are physically stored."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "Most existing tools address only one side of this problem. Circuit "
                    "design environments usually assume parts exist. Inventory tools may "
                    "track stock, but rarely guide physical retrieval. General-purpose AI "
                    "can propose circuits, but often ignores owned-stock constraints, exact "
                    "variants, and the practical step of finding the required parts in "
                    "drawers, bins, or shelves."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "AURA is designed so that a reader can understand both contributions at "
                    "once: it is an inventory-aware AI-assisted circuit creation system and "
                    "a host-and-node localization system for real component retrieval. The "
                    "core idea is that both capabilities should operate on the same inventory "
                    "substrate, so that accepted build suggestions can immediately transition "
                    "into physically actionable locate operations."
                ),
            },
            {
                "type": "subheading",
                "text": "A. Problem Scope",
            },
            {
                "type": "paragraph",
                "text": (
                    "AURA therefore answers two practical questions together: what can be "
                    "built from the exact owned parts that are available, and where are the "
                    "required parts physically located. Build feasibility depends on exact "
                    "values, quantities, and variants. Retrieval depends on known storage "
                    "addresses and reliable physical highlighting. Solving only one side "
                    "still leaves the user blocked by the other."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Contributions",
            },
            {
                "type": "paragraph",
                "text": (
                    "The main contribution of AURA is a complete prototype that unifies "
                    "inventory-aware AI-assisted circuit creation with addressed physical "
                    "component localization. The specific contributions are: 1) a structured "
                    "circuit assistance layer that produces reviewable, inventory-aware "
                    "candidate circuits; 2) a host-and-node cyber-physical localization "
                    "prototype for exact part retrieval; 3) a shared inventory model that "
                    "captures exact values, quantities, and storage addresses; and 4) an "
                    "integrated workflow in which accepted circuit outputs lead directly to "
                    "retrieval actions in physical space."
                ),
            },
        ],
    },
    {
        "title": "II. DUAL-CORE SYSTEM ARCHITECTURE",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "AURA is organized around two equally important operational cores that "
                    "share one inventory substrate. The first core is an inventory-aware "
                    "circuit creation path that uses deterministic structure and local rules "
                    "to keep AI assistance reviewable. The second core is a host-guided "
                    "localization path that transforms exact-part knowledge into addressed "
                    "physical highlighting through distributed nodes."
                ),
            },
            {
                "type": "figure",
                "path": ARCHITECTURE_FIG,
                "caption": (
                    "Fig. 1. Integrated AURA architecture linking inventory-aware circuit "
                    "creation to host-guided physical localization."
                ),
                "width_in": 3.2,
            },
            {
                "type": "subheading",
                "text": "A. Inventory-Aware AI-Assisted Circuit Creation Core",
            },
            {
                "type": "paragraph",
                "text": (
                    "The circuit-creation core accepts user intent, inventory context, and "
                    "structured constraints. Instead of treating AI output as free-form chat, "
                    "AURA keeps candidate designs tied to structured circuit state and local "
                    "deterministic checks. This makes build suggestions inspectable, easier "
                    "to review, and grounded in owned stock rather than idealized parts."
                ),
            },
            {
                "type": "figure",
                "path": ROOT / "ppt_assets" / "conference_media" / "conference_workspace.png",
                "caption": (
                    "Fig. 2. Structured support workspace used for inventory-aware circuit "
                    "creation, review, and deterministic assistance."
                ),
                "width_in": 3.15,
            },
            {
                "type": "subheading",
                "text": "B. Host-Guided Component Localization Core",
            },
            {
                "type": "paragraph",
                "text": (
                    "The localization core is driven by an ESP32-class host with a compact "
                    "local display and joystick interface. The host is responsible for part "
                    "lookup, stock review, node addressing, and locate activation. Its role "
                    "is intentionally independent enough that physical retrieval does not "
                    "require a smartphone in normal use."
                ),
            },
            {
                "type": "figure",
                "path": ROOT / "ppt_assets" / "conference_media" / "image13.png",
                "caption": (
                    "Fig. 3. AURA host prototype used for local lookup, stock review, and "
                    "locate activation."
                ),
                "width_in": 3.0,
            },
            {
                "type": "paragraph",
                "text": (
                    "Compact low-power locator nodes based on the ATtiny404 and nRF24L01 "
                    "provide the distributed actuation layer. With WS2812-based visual "
                    "indication and optional extensibility to other cues, nodes map digital "
                    "part selection to physical storage positions. This allows the system to "
                    "highlight the exact drawer, strip, or zone containing the accepted part."
                ),
            },
            {
                "type": "figure",
                "path": ROOT / "ppt_assets" / "conference_media" / "image14.png",
                "caption": (
                    "Fig. 4. Locator node prototype used for addressed physical highlighting "
                    "of stored components."
                ),
                "width_in": 2.6,
            },
            {
                "type": "subheading",
                "text": "C. Shared Inventory and Address Layer",
            },
            {
                "type": "paragraph",
                "text": (
                    "Both cores depend on the same inventory layer. Records are intended to "
                    "capture component family, exact value or variant, quantity, and physical "
                    "storage address. This shared substrate is what makes AURA more than a "
                    "pair of disconnected tools. The same inventory knowledge constrains "
                    "candidate circuits and powers physical retrieval."
                ),
            },
        ],
    },
    {
        "title": "III. INTEGRATED USER WORKFLOW",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "The practical effect of AURA is best understood through the transition "
                    "from build intent to physical assembly. A user can begin with an idea, "
                    "a function, or a component need. AURA then supports circuit creation "
                    "and retrieval as one continuous flow rather than leaving them in "
                    "separate tools."
                ),
            },
            {
                "type": "subheading",
                "text": "A. From Build Intent to Candidate Circuit",
            },
            {
                "type": "paragraph",
                "text": (
                    "AURA's assistance layer can propose candidate circuits that remain tied "
                    "to inventory context and structured review. This is important because "
                    "users do not merely want any plausible circuit. They need one that is "
                    "compatible with what they actually own and can inspect before accepting."
                ),
            },
            {
                "type": "subheading",
                "text": "B. From Accepted Circuit to Exact Part Set",
            },
            {
                "type": "paragraph",
                "text": (
                    "Once a candidate circuit is accepted, AURA resolves the required part "
                    "set against the shared inventory layer. This stage translates the "
                    "design-side result into exact part identities, exact values, and "
                    "required quantities. The workflow therefore remains grounded in what is "
                    "buildable, not merely what is theoretically valid."
                ),
            },
            {
                "type": "subheading",
                "text": "C. From Exact Part Set to Physical Retrieval",
            },
            {
                "type": "paragraph",
                "text": (
                    "After the exact part set is confirmed, the host resolves storage "
                    "addresses and issues addressed commands to the mapped locator nodes. "
                    "The relevant outputs then highlight the physical positions associated "
                    "with the accepted circuit or requested part. This closes the loop "
                    "between design and assembly in a way that software-only systems do not."
                ),
            },
            {
                "type": "subheading",
                "text": "D. Host-First Operation with Optional Smartphone Support",
            },
            {
                "type": "paragraph",
                "text": (
                    "A smartphone or richer software surface can support tasks such as bulk "
                    "inventory input, richer naming, and more comfortable project browsing. "
                    "However, the host remains the primary embedded retrieval interface. "
                    "This keeps AURA positioned as a real hardware product with software "
                    "support rather than a phone application with attached electronics."
                ),
            },
        ],
    },
    {
        "title": "IV. CURRENT PROTOTYPE REALIZATION AND EFFECT",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "The current AURA realization is intended as a complete prototype, not "
                    "just a concept sketch or isolated demo. The prototype already spans "
                    "structured circuit assistance, inventory-aware filtering, host-driven "
                    "lookup, node-addressed localization, and hardware-guided retrieval."
                ),
            },
            {
                "type": "subheading",
                "text": "A. Implemented Prototype Capabilities",
            },
            {
                "type": "paragraph",
                "text": (
                    "The integrated prototype direction supports inventory-constrained "
                    "circuit proposals, structured review of candidate builds, single-part "
                    "lookup, exact-value stock checking, addressed locate activation, and "
                    "multi-part retrieval concepts for accepted builds. These capabilities "
                    "matter because they demonstrate that circuit creation and physical "
                    "localization are not being claimed independently, but are being "
                    "connected through one operational system."
                ),
            },
            {
                "type": "subheading",
                "text": "B. Practical Effect on User Workflow",
            },
            {
                "type": "paragraph",
                "text": (
                    "On the design side, AURA reduces the gap between a vague build idea and "
                    "a concrete candidate circuit grounded in owned stock. On the retrieval "
                    "side, it reduces the drawer-by-drawer or shelf-by-shelf search burden "
                    "that normally follows once the needed parts are identified. The result "
                    "is a shorter, more inspectable path from intent to assembly."
                ),
            },
            {
                "type": "subheading",
                "text": "C. Prototype Positioning",
            },
            {
                "type": "paragraph",
                "text": (
                    "The novelty of AURA should be understood in the integration itself. "
                    "The paper does not rely on an unsupported claim that each isolated "
                    "sub-part is globally unprecedented. Instead, it presents a unified "
                    "cyber-physical prototype in which inventory-aware circuit assistance "
                    "and addressed localization reinforce one another through shared system "
                    "state."
                ),
            },
        ],
    },
    {
        "title": "V. DISCUSSION",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "AURA should not be evaluated as a claim of universal circuit "
                    "correctness or full simulator coverage. Its value lies in combining two "
                    "normally disconnected forms of help: assisting users in identifying a "
                    "buildable circuit from owned stock, and then guiding them to the real "
                    "components required to assemble it."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "This balanced framing is important. If the paper overemphasizes only "
                    "the localization side, the assistance layer appears secondary. If it "
                    "overemphasizes only the circuit-creation side, the system risks being "
                    "misread as another software tool. The correct interpretation is that "
                    "AURA is a dual-core cyber-physical system whose practical impact comes "
                    "from joining both functions through one inventory-aware workflow."
                ),
            },
        ],
    },
    {
        "title": "VI. CONCLUSION",
        "blocks": [
            {
                "type": "paragraph",
                "text": (
                    "AURA presents a complete prototype in which inventory-aware AI-assisted "
                    "circuit creation and distributed component localization operate as one "
                    "integrated electronics assistance system. By grounding candidate "
                    "circuits in owned stock and then connecting accepted results to "
                    "addressed physical retrieval, the prototype reduces friction on both "
                    "the design side and the assembly side of electronics work."
                ),
            },
            {
                "type": "paragraph",
                "text": (
                    "Future work will strengthen validation, broaden storage coverage, refine "
                    "the embedded hardware, and expand support surfaces without losing the "
                    "core identity established here. The central result remains stable: "
                    "AURA treats circuit creation and physical retrieval as two parts of the "
                    "same problem and implements them together in one cyber-physical "
                    "prototype."
                ),
            },
        ],
    },
]

REFERENCES = [
    '[1] Espressif Systems, "ESP32 Technical Reference Manual," 2023.',
    '[2] Microchip Technology Inc., "tinyAVR 0-series Family Data Sheet," 2024.',
    '[3] Nordic Semiconductor, "nRF24L01+ Single Chip 2.4 GHz Transceiver Product Specification," 2019.',
    '[4] Worldsemi, "WS2812B Intelligent Control LED Integrated Light Source Data Sheet," 2020.',
    '[5] Bluetooth SIG, "Bluetooth Core Specification," Version 5.x.',
    '[6] Arduino, "SPI Communication," Arduino Documentation.',
    '[7] SQLite, "SQLite Documentation," Documentation.',
    '[8] P. Horowitz and W. Hill, The Art of Electronics, 3rd ed. Cambridge, U.K.: Cambridge Univ. Press, 2015.',
    '[9] IETF, "JavaScript Object Notation (JSON) Patch," RFC 6902, 2013.',
]


def create_architecture_figure():
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    def box(x1, y1, x2, y2, text):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, outline="black", width=3, fill="white")
        words = text.split()
        lines = []
        cur = []
        for word in words:
            trial = " ".join(cur + [word])
            if draw.textlength(trial, font=font) > (x2 - x1 - 30) and cur:
                lines.append(" ".join(cur))
                cur = [word]
            else:
                cur.append(word)
        if cur:
            lines.append(" ".join(cur))
        total_h = len(lines) * 16
        y = y1 + ((y2 - y1 - total_h) / 2)
        for line in lines:
            w = draw.textlength(line, font=font)
            draw.text((x1 + ((x2 - x1 - w) / 2), y), line, fill="black", font=font)
            y += 16

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="black", width=3)
        if x1 == x2:
            draw.polygon([(x2, y2), (x2 - 8, y2 - 16), (x2 + 8, y2 - 16)], fill="black")
        else:
            draw.polygon([(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)], fill="black")

    box(560, 40, 940, 120, "User intent / build need")
    box(120, 200, 520, 320, "Inventory-aware AI-assisted circuit creation")
    box(580, 200, 920, 320, "Shared inventory and exact-part records")
    box(980, 200, 1380, 320, "Host-guided physical localization")
    box(120, 420, 520, 520, "Structured candidate circuit and review")
    box(580, 420, 920, 520, "Accepted part set and quantity fit")
    box(980, 420, 1380, 520, "Addressed node commands")
    box(980, 640, 1380, 760, "Distributed nodes and physical retrieval")
    box(120, 640, 520, 760, "Phone / workspace support surface")

    arrow(750, 120, 320, 200)
    arrow(750, 120, 750, 200)
    arrow(750, 120, 1180, 200)
    arrow(320, 320, 320, 420)
    arrow(750, 320, 750, 420)
    arrow(1180, 320, 1180, 420)
    arrow(520, 470, 580, 470)
    arrow(920, 470, 980, 470)
    arrow(1180, 520, 1180, 640)
    arrow(320, 700, 320, 520)

    draw.text((520, 385), "inventory constrains both cores", fill="black", font=font)
    draw.text((1040, 560), "accepted outputs become locate actions", fill="black", font=font)
    draw.text((70, 825), "AURA integrates circuit creation and localization through one inventory substrate.", fill="black", font=font)

    img.save(ARCHITECTURE_FIG)


def set_page_layout(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.63)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def set_two_columns(section, num=2, space_twips=360):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    if cols.getparent() is None:
        sect_pr.append(cols)


def set_run_font(run, size_pt, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph_text(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.16):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Inches(first_indent)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size)
    return p


def add_center_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 10, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 10, bold=True)
    return p


def add_abstract_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(f"{label}\u2014")
    set_run_font(label_run, 9, bold=True, italic=True)
    body_run = p.add_run(text)
    set_run_font(body_run, 9)
    return p


def add_figure(doc, path, caption, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    set_run_font(run, 8, italic=True)


def add_references(doc):
    add_center_heading(doc, "REFERENCES")
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(ref)
        set_run_font(run, 8)


def build_markdown():
    lines = [
        f"# {TITLE}",
        "",
        *AUTHOR_LINES,
        "",
        "## Abstract",
        "",
        ABSTRACT,
        "",
        "## Index Terms",
        "",
        INDEX_TERMS,
        "",
    ]

    for section in SECTIONS:
        lines.append(f"## {section['title']}")
        lines.append("")
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                lines.append(block["text"])
                lines.append("")
            elif block["type"] == "subheading":
                lines.append(f"### {block['text']}")
                lines.append("")
            elif block["type"] == "figure":
                lines.append(f"![{block['caption']}]({block['path'].as_posix()})")
                lines.append("")
                lines.append(f"*{block['caption']}*")
                lines.append("")

    lines.append("## REFERENCES")
    lines.append("")
    lines.extend(REFERENCES)
    lines.append("")
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_docx():
    doc = Document()
    set_page_layout(doc.sections[0])

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(TITLE)
    set_run_font(run, 18, bold=True)

    for line in AUTHOR_LINES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, 10)

    doc.add_paragraph()
    add_abstract_paragraph(doc, "Abstract", ABSTRACT)
    add_abstract_paragraph(doc, "Index Terms", INDEX_TERMS)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page_layout(body_section)
    set_two_columns(body_section)

    for section in SECTIONS:
        add_center_heading(doc, section["title"])
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                add_paragraph_text(doc, block["text"])
            elif block["type"] == "subheading":
                add_subheading(doc, block["text"])
            elif block["type"] == "figure":
                add_figure(doc, block["path"], block["caption"], block["width_in"])

    add_references(doc)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    create_architecture_figure()
    build_markdown()
    build_docx()
    print(f"Wrote {ARCHITECTURE_FIG.name}")
    print(f"Wrote {OUTPUT_MD.name}")
    print(f"Wrote {OUTPUT_DOCX.name}")
